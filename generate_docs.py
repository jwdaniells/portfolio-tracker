"""
generate_docs.py
Generates the Portfolio Tracker System Documentation as a Word document.
Run: python generate_docs.py
Output: Portfolio_Tracker_System_Documentation.docx
"""

import io, os, textwrap
from pathlib import Path
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import matplotlib.patheffects as pe
from PIL import Image

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIR = Path(__file__).parent
TODAY = date.today().strftime("%-d %B %Y")

# ── Colour palette ─────────────────────────────────────────────────────────────
GOLD   = RGBColor(0xC9, 0xA8, 0x4C)
DARK   = RGBColor(0x0F, 0x19, 0x23)
MID    = RGBColor(0x12, 0x1E, 0x2B)
LIGHT  = RGBColor(0xE8, 0xDC, 0xC8)
GREY   = RGBColor(0x6A, 0x7D, 0x8F)
GREEN  = RGBColor(0x70, 0xAD, 0x47)
RED    = RGBColor(0xE0, 0x70, 0x60)
BLUE   = RGBColor(0x44, 0x72, 0xC4)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARKBLUE = RGBColor(0x1A, 0x2E, 0x4A)  # dark navy for body text

# Hex versions for matplotlib
BG      = "#0f1923"
CARD    = "#121e2b"
GoldH   = "#c9a84c"
BlueH   = "#4472C4"
GreenH  = "#70AD47"
RedH    = "#e07060"
GreyH   = "#6a7d8f"
LightH  = "#e8dcc8"

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour.lstrip("#"))
    tcPr.append(shd)

def set_row_shading(row, hex_colour):
    for cell in row.cells:
        set_cell_bg(cell, hex_colour)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A84C")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.font.size   = Pt(13)
    run.font.bold   = True
    run.font.color.rgb = GOLD
    run.font.name   = "Calibri"
    add_horizontal_rule(doc)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = DARKBLUE
    run.font.name  = "Calibri"

def body(doc, text, italic=False, colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.italic = italic
    run.font.name   = "Calibri"
    run.font.color.rgb = colour if colour else DARKBLUE

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.name  = "Calibri"
    run.font.color.rgb = DARKBLUE

def code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = GOLD

def fig_to_docx(doc, fig, width_cm=16):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))
    plt.close(fig)

def add_artifact_table(doc, rows):
    """rows: list of (filename, type, purpose, notes)"""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(["File / Artefact", "Type", "Purpose", "Notes"]):
        cell = hdr.cells[i]
        set_cell_bg(cell, "0a1420")
        run = cell.paragraphs[0].add_run(h.upper())
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = GOLD
        run.font.name = "Calibri"
    widths = [Cm(4.2), Cm(2.4), Cm(6.5), Cm(4.0)]
    for i, col in enumerate(tbl.columns):
        for cell in col.cells:
            cell.width = widths[i]
    for idx, (fname, ftype, purpose, notes) in enumerate(rows):
        row = tbl.add_row()
        bg = "121e2b" if idx % 2 == 0 else "0f1923"
        set_row_shading(row, bg)
        data = [fname, ftype, purpose, notes]
        for i, val in enumerate(data):
            cell = row.cells[i]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = "Courier New" if i == 0 else "Calibri"
            run.font.color.rgb = GOLD if i == 0 else LIGHT

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════

def make_architecture_diagram():
    fig, ax = plt.subplots(figsize=(14, 7), facecolor=BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 7)
    ax.axis("off")

    def box(x, y, w, h, label, sub="", color=BlueH, text_color=LightH):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
            linewidth=1.5, edgecolor=color, facecolor=color + "22")
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.12 if sub else 0), label,
            ha="center", va="center", color=text_color,
            fontsize=9, fontweight="bold", fontfamily="monospace")
        if sub:
            ax.text(x + w/2, y + h/2 - 0.22, sub,
                ha="center", va="center", color=GreyH, fontsize=7.5)

    def arrow(x1, y1, x2, y2, label="", col=GreyH):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", color=col, lw=1.5))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx, my + 0.12, label, ha="center", va="bottom",
                    color=col, fontsize=7.5, style="italic")

    # GitHub
    box(5.5, 5.5, 3, 1.0, "GitHub Repository", "jwdaniells/portfolio-tracker", GoldH)
    # GitHub Actions
    box(1.0, 5.5, 3.0, 1.0, "GitHub Actions", "Scheduled 16:45 UTC weekdays", GreenH)
    # GitHub Pages
    box(10.0, 5.5, 3.0, 1.0, "GitHub Pages", "jwdaniells.github.io/...", BlueH)
    # Yahoo Finance
    box(1.0, 3.5, 3.0, 1.0, "Yahoo Finance API", "Price data (LSE + Crypto)", GreyH)
    # Email
    box(1.0, 1.5, 3.0, 1.0, "Yahoo Mail SMTP", "jdaniells@yahoo.co.uk", GoldH)
    # User browser
    box(10.0, 3.5, 3.0, 1.0, "User's Browser", "Password gate → App", BlueH)
    # Local Mac
    box(10.0, 1.5, 3.0, 1.0, "Local Mac", "pull_updates.command", GreyH)
    # JSON data
    box(5.5, 3.5, 3.0, 1.0, "JSON Data Files", "price_history-2.json\nanalysis.json", GoldH)

    # Arrows
    arrow(4.0, 6.0, 5.5, 6.0, "trigger / commit")
    arrow(4.0, 4.0, 5.5, 4.0, "fetch prices")
    arrow(4.0, 2.0, 4.0, 3.5, "sends email")
    arrow(8.5, 6.0, 10.0, 6.0, "auto-deploy")
    arrow(8.5, 4.0, 10.0, 4.0, "serves files")
    arrow(11.5, 3.5, 11.5, 2.5, "git pull")
    arrow(5.5, 4.0, 4.0, 4.0, "")
    arrow(7.0, 5.5, 7.0, 4.5, "commits JSON")

    ax.set_title("Portfolio Tracker — System Architecture", color=LightH,
                 fontsize=12, fontweight="bold", pad=10)
    fig.tight_layout(pad=0.5)
    return fig


def make_daily_flow_diagram():
    fig, ax = plt.subplots(figsize=(14, 5), facecolor=BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 5)
    ax.axis("off")

    steps = [
        (1.0,  "1. Schedule\nTrigger", "16:45 UTC\nweekdays", GoldH),
        (3.2,  "2. Fetch\nPrices", "Yahoo Finance\nAPI", BlueH),
        (5.4,  "3. Update\nJSON", "price_history\n-2.json", GreenH),
        (7.6,  "4. Git\nCommit", "Auto-commit\nto main", GoldH),
        (9.8,  "5. GitHub\nPages Deploy", "Auto rebuild\n~1 min", BlueH),
        (12.0, "6. Send\nEmail", "Stats to\njdaniells@...", GreenH),
    ]

    for x, label, sub, col in steps:
        rect = mpatches.FancyBboxPatch((x, 1.8), 1.8, 1.4,
            boxstyle="round,pad=0.1", linewidth=1.5,
            edgecolor=col, facecolor=col + "22")
        ax.add_patch(rect)
        ax.text(x + 0.9, 2.8, label, ha="center", va="center",
                color=LightH, fontsize=8.5, fontweight="bold")
        ax.text(x + 0.9, 2.2, sub, ha="center", va="center",
                color=GreyH, fontsize=7.5)

    for i in range(len(steps) - 1):
        x = steps[i][0] + 1.8
        ax.annotate("", xy=(steps[i+1][0], 2.5), xytext=(x, 2.5),
            arrowprops=dict(arrowstyle="->", color=GoldH, lw=1.5))

    # Error path
    ax.annotate("", xy=(3.2, 1.5), xytext=(5.4, 1.5),
        arrowprops=dict(arrowstyle="<-", color=RedH, lw=1, linestyle="dashed"))
    ax.text(4.3, 1.25, "On error: abort + failure email", ha="center",
            color=RedH, fontsize=7.5, style="italic")

    ax.set_title("Daily Automated Price Fetch — Process Flow", color=LightH,
                 fontsize=11, fontweight="bold", pad=8)
    fig.tight_layout(pad=0.5)
    return fig


def make_user_access_flow():
    fig, ax = plt.subplots(figsize=(14, 4.5), facecolor=BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 4.5)
    ax.axis("off")

    def box(x, y, w, h, text, col, shape="rect"):
        if shape == "diamond":
            diamond = plt.Polygon(
                [(x+w/2, y+h), (x+w, y+h/2), (x+w/2, y), (x, y+h/2)],
                fill=True, facecolor=col+"22", edgecolor=col, linewidth=1.5)
            ax.add_patch(diamond)
            ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                    color=LightH, fontsize=8, fontweight="bold")
        else:
            rect = mpatches.FancyBboxPatch((x, y), w, h,
                boxstyle="round,pad=0.08", linewidth=1.5,
                edgecolor=col, facecolor=col+"22")
            ax.add_patch(rect)
            ax.text(x+w/2, y+h/2, text, ha="center", va="center",
                    color=LightH, fontsize=8, fontweight="bold")

    def arr(x1, y1, x2, y2, label="", col=GreyH):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", color=col, lw=1.5))
        if label:
            ax.text((x1+x2)/2, (y1+y2)/2 + 0.1, label, ha="center",
                    color=col, fontsize=7.5, style="italic")

    box(0.3,  2.8, 2.0, 0.9, "Open URL\n(GitHub Pages)", BlueH)
    box(3.0,  2.8, 2.0, 0.9, "Login Screen\nDisplayed", GoldH)
    box(5.7,  2.8, 2.2, 0.9, "Password\nCorrect?", GoldH, "diamond")
    box(9.2,  2.8, 2.0, 0.9, "App Loads\n(JSX rendered)", GreenH)
    box(11.5, 2.8, 2.0, 0.9, "View Dashboard\n/ Tabs", GreenH)
    box(5.7,  1.0, 2.2, 0.9, "Error Message\n3s then clear", RedH)

    arr(2.3, 3.25, 3.0, 3.25)
    arr(5.0, 3.25, 5.7, 3.25)
    arr(7.9, 3.25, 9.2, 3.25, "Yes", GreenH)
    arr(13.5, 3.25, 13.5, 3.25)
    arr(11.2, 3.25, 11.5, 3.25)
    arr(6.8, 2.8, 6.8, 1.9, "No", RedH)
    arr(6.8, 1.0, 5.0, 3.25, "Retry", GoldH)

    ax.text(7.0, 0.4, "Session persisted in sessionStorage — no re-prompt on refresh within same browser session",
            ha="center", color=GreyH, fontsize=7.5, style="italic")

    ax.set_title("User Access Flow — Password Gate", color=LightH,
                 fontsize=11, fontweight="bold", pad=8)
    fig.tight_layout(pad=0.5)
    return fig


def make_data_flow_diagram():
    fig, ax = plt.subplots(figsize=(14, 5.5), facecolor=BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 14); ax.set_ylim(0, 5.5)
    ax.axis("off")

    def box(x, y, w, h, title, sub="", col=BlueH):
        rect = mpatches.FancyBboxPatch((x, y), w, h,
            boxstyle="round,pad=0.1", linewidth=1.5,
            edgecolor=col, facecolor=col+"18")
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2+(0.15 if sub else 0), title, ha="center",
                va="center", color=LightH, fontsize=8.5, fontweight="bold")
        if sub:
            ax.text(x+w/2, y+h/2-0.22, sub, ha="center", va="center",
                    color=GreyH, fontsize=7.5)

    def arr(x1, y1, x2, y2, label="", col=GreyH, double=False):
        style = "<->" if double else "->"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle=style, color=col, lw=1.4))
        if label:
            ax.text((x1+x2)/2, (y1+y2)/2+0.12, label, ha="center",
                    color=col, fontsize=7.5, style="italic")

    # Data store
    box(5.5, 2.8, 3.0, 1.2, "price_history-2.json", "Master data store", GoldH)
    # fetch_prices.py
    box(0.5, 3.8, 2.8, 0.9, "fetch_prices.py", "GitHub Actions job", GreenH)
    # analysis.json
    box(5.5, 0.8, 3.0, 1.0, "analysis.json", "AI analysis output", GreyH)
    # reminders.json
    box(5.5, 4.8, 3.0, 0.6, "reminders.json", "", GreyH)
    # React App
    box(10.0, 2.8, 3.2, 1.2, "portfolio_tracker\n_v1.20.jsx", "React SPA", BlueH)
    # retirement
    box(10.0, 0.8, 3.2, 1.0, "retirement_planner\n.jsx", "Retirement calc", BlueH)
    # index.html
    box(10.0, 4.5, 3.2, 0.8, "index.html", "Entry / auth gate", GoldH)

    arr(3.3, 4.25, 5.5, 3.4, "writes prices", GreenH)
    arr(8.5, 3.4, 10.0, 3.4, "fetches via HTTP", BlueH, double=True)
    arr(8.5, 1.3, 10.0, 1.3, "fetches via HTTP", BlueH)
    arr(8.5, 1.3, 10.0, 1.3)
    arr(7.0, 2.8, 7.0, 1.8, "read", GreyH)
    arr(11.6, 4.5, 11.6, 4.0, "loads JSX", GoldH)
    arr(7.0, 4.8, 8.05, 4.3, "read", GreyH)

    ax.set_title("Data Architecture & File Relationships", color=LightH,
                 fontsize=11, fontweight="bold", pad=8)
    fig.tight_layout(pad=0.5)
    return fig


# ══════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover Page ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DANIELLS PORTFOLIO & RETIREMENT TRACKER")
    run.font.size  = Pt(24)
    run.font.bold  = True
    run.font.color.rgb = GOLD
    run.font.name  = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("System Documentation")
    run.font.size  = Pt(16)
    run.font.color.rgb = LIGHT
    run.font.name  = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─────────────────────────────────────")
    run.font.color.rgb = GOLD

    doc.add_paragraph()

    for line, size, col in [
        (f"Version: 1.20", 11, DARKBLUE),
        (f"Date: {TODAY}", 11, DARKBLUE),
        ("Author: John Daniells", 11, DARKBLUE),
        ("Classification: Personal / Confidential", 10, GREY),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.color.rgb = col
        run.font.name = "Calibri"

    doc.add_page_break()

    # ── 1. Purpose & Intent ────────────────────────────────────────────────────
    heading1(doc, "1. Purpose & Intent")
    body(doc, (
        "The Daniells Portfolio Tracker is a personal financial management application designed "
        "to provide a real-time, consolidated view of an investment portfolio spanning multiple "
        "accounts, asset classes, and wrappers. It serves as a single source of truth for "
        "portfolio performance, allocation, goal progress, and retirement planning."
    ))
    heading2(doc, "Primary Objectives")
    for b in [
        "Consolidate holdings across ISAs, SIPPs, GIAs, and crypto wallets into one view.",
        "Automate daily price fetching from Yahoo Finance after LSE market close (16:30 GMT).",
        "Track progress toward a £1,400,000 retirement goal by February 2031.",
        "Provide a browser-accessible interface via GitHub Pages — no local server required.",
        "Send an automated daily email summary with portfolio value and per-account movements.",
        "Maintain a complete price history for charting, trend analysis, and audit purposes.",
    ]:
        bullet(doc, b)

    heading2(doc, "Design Principles")
    for b in [
        "Static-first: No backend server. All logic runs in the browser via React + Babel.",
        "Zero-install: Runs directly from any browser — desktop, tablet, or mobile.",
        "Secure by default: Password-gated with SHA-256 hashing; data only loads post-authentication.",
        "Automated: GitHub Actions fetches prices and publishes updates without manual intervention.",
        "Portable: Entire system lives in a single GitHub repository.",
    ]:
        bullet(doc, b)

    # ── 2. System Architecture ─────────────────────────────────────────────────
    heading1(doc, "2. System Architecture")
    body(doc, (
        "The system is composed of five logical layers: a data store (JSON files), an automation "
        "layer (GitHub Actions), a hosting layer (GitHub Pages), a presentation layer (React SPA), "
        "and a notification layer (Yahoo Mail SMTP). These are all orchestrated via a single "
        "GitHub repository."
    ))
    fig = make_architecture_diagram()
    fig_to_docx(doc, fig, width_cm=16)
    body(doc, "Figure 1: High-level system architecture showing all components and their relationships.", italic=True, colour=GREY)

    heading2(doc, "Hosting & Deployment")
    for b in [
        "Repository: github.com/jwdaniells/portfolio-tracker (public)",
        "GitHub Pages URL: https://jwdaniells.github.io/portfolio-tracker/",
        "Branch: main — all commits to main trigger an automatic Pages rebuild (~1 minute).",
        "No build pipeline required — files are served as-is (HTML, JSX, JSON).",
    ]:
        bullet(doc, b)

    heading2(doc, "Technology Stack")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Layer", "Technology", "Details"]):
        set_cell_bg(hdr.cells[i], "0a1420")
        run = hdr.cells[i].paragraphs[0].add_run(h.upper())
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = GOLD; run.font.name = "Calibri"
    for idx, (layer, tech, detail) in enumerate([
        ("Frontend", "React 18 (UMD)", "No build step — loaded from unpkg CDN"),
        ("Transpiler", "Babel Standalone", "JSX-to-JS transform in the browser at runtime"),
        ("Styling", "Inline CSS (JS objects)", "Dark theme, no external CSS framework"),
        ("Charts", "SVG (hand-built)", "Custom SVG paths, no third-party chart library"),
        ("Data", "JSON files", "price_history-2.json, analysis.json, reminders.json"),
        ("Automation", "GitHub Actions", "Python 3.11 on Ubuntu, runs on cron schedule"),
        ("Hosting", "GitHub Pages", "Served from main branch root directory"),
        ("Notifications", "SMTP / action-send-mail", "Yahoo Mail SMTP, credentials in GitHub Secrets"),
        ("Price source", "Yahoo Finance v8 API", "Direct HTTP — no yfinance library dependency"),
        ("Security", "SHA-256 password gate", "browser crypto.subtle — client-side login screen"),
    ]):
        row = tbl.add_row()
        set_row_shading(row, "121e2b" if idx % 2 == 0 else "0f1923")
        for i, val in enumerate([layer, tech, detail]):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = LIGHT

    # ── 3. Component Inventory ─────────────────────────────────────────────────
    heading1(doc, "3. Component Inventory")
    body(doc, "All files in the repository and their roles:")

    add_artifact_table(doc, [
        ("index.html",                  "HTML",         "Application entry point. Ships the React CDN imports, password gate logic and JSX loader.",
                                                         "Modified to add SHA-256 auth gate (Mar 2026)"),
        ("portfolio_tracker_v1.20.jsx", "React/JSX",    "Main single-page application. All tabs: Dashboard, Accounts, Allocation, History, Goal, Analysis, Crypto.",
                                                         "7 tabs; ~500 lines; no build required"),
        ("retirement.html",             "HTML",         "Separate entry point for the Retirement Planner tool.",
                                                         "Linked from nav bar in main app"),
        ("retirement_planner.jsx",      "React/JSX",    "Retirement income and drawdown planning calculator.",
                                                         "State pension, drawdown, ISA/SIPP modelling"),
        ("price_history-2.json",        "JSON",         "Master data store. Holds all account definitions, holdings, prices, history, crypto, and metadata.",
                                                         "Updated daily by GitHub Actions"),
        ("analysis.json",               "JSON",         "AI-generated portfolio analysis and commentary. Read by the Analysis tab.",
                                                         "Manually updated; non-blocking if absent"),
        ("reminders.json",              "JSON",         "User-defined reminders and notes displayed in the app.",
                                                         "Manually maintained"),
        ("fetch_prices.py",             "Python",       "Price fetching script. Calls Yahoo Finance v8 API for all equity, fund, and crypto tickers. Updates JSON.",
                                                         "Run by GitHub Actions at 16:45 UTC weekdays"),
        (".github/workflows/fetch-prices.yml", "YAML",  "GitHub Actions workflow definition. Schedules and orchestrates the fetch, commit, and email steps.",
                                                         "Cron: 45 16 * * 1-5"),
        ("pull_updates.command",        "Shell",        "Double-click script for macOS. Runs git pull to sync latest prices from GitHub to local machine.",
                                                         "Optional — only needed for local viewing"),
        (".gitignore",                  "Config",       "Excludes local-only files from the repository (fetch_analysis.py, *.command etc.)",
                                                         "pull_updates.command IS tracked"),
    ])

    # ── 4. Process Flows ────────────────────────────────────────────────────────
    heading1(doc, "4. Process Flows")

    heading2(doc, "4.1  Daily Automated Price Fetch")
    fig = make_daily_flow_diagram()
    fig_to_docx(doc, fig, width_cm=16)
    body(doc, "Figure 2: Daily automated price fetch and publish flow.", italic=True, colour=GREY)

    body(doc, "The cron schedule is: 45 16 * * 1-5 (16:45 UTC, Monday–Friday). This is 15 minutes after the London Stock Exchange closes at 16:30 GMT. In BST (summer, UTC+1) the run time is 17:45 local.")
    for b in [
        "GitHub Actions checks out the repository on an Ubuntu runner.",
        "Python 3.11 is installed; requests library installed via pip.",
        "fetch_prices.py calls the Yahoo Finance v8 chart API for each ticker.",
        "UK equity and fund prices are returned in pence (GBp) and automatically divided by 100.",
        "Crypto prices (BTC, ETH, XRP, ADA) are returned directly in GBP.",
        "If zero prices are fetched, the script aborts — the JSON is never written with corrupt data.",
        "Updated price_history-2.json is committed back to the repo with message 'Auto-update prices DD Mon YYYY'.",
        "GitHub Pages automatically rebuilds from the new commit within ~1 minute.",
        "An email is sent to jdaniells@yahoo.co.uk with per-account stats and a link to the tracker.",
        "On any failure, a failure alert email is sent instead.",
    ]:
        bullet(doc, b)

    heading2(doc, "4.2  User Access Flow")
    fig = make_user_access_flow()
    fig_to_docx(doc, fig, width_cm=16)
    body(doc, "Figure 3: User authentication and application load flow.", italic=True, colour=GREY)

    body(doc, "The password gate operates entirely client-side:")
    for b in [
        "The login screen is rendered as plain HTML/CSS — no React/JSX is loaded at this stage.",
        "The user enters their password; it is hashed in-browser using window.crypto.subtle (SHA-256).",
        "The hash is compared against a hardcoded hash constant in index.html. The plaintext password is never stored.",
        "On success, sessionStorage is flagged and the JSX app is fetched, transpiled and mounted.",
        "Session persists for the duration of the browser tab — refreshing does not require re-entry.",
        "If the browser tab is closed, the session flag is cleared and the password must be re-entered.",
    ]:
        bullet(doc, b)

    heading2(doc, "4.3  Data Flow & File Relationships")
    fig = make_data_flow_diagram()
    fig_to_docx(doc, fig, width_cm=16)
    body(doc, "Figure 4: Data architecture and relationships between files.", italic=True, colour=GREY)

    heading2(doc, "4.4  Local Sync Flow")
    body(doc, "When the user wants to view updated prices locally (e.g. via localhost:8765):")
    for b in [
        "Double-click pull_updates.command in Finder.",
        "Script runs: git pull in the project directory.",
        "Latest price_history-2.json is downloaded from GitHub.",
        "Refresh the local browser to see updated data.",
    ]:
        bullet(doc, b)

    # ── 5. Data Architecture ───────────────────────────────────────────────────
    heading1(doc, "5. Data Architecture")
    heading2(doc, "price_history-2.json Structure")
    body(doc, "This is the master data file. Its top-level structure is:")
    for field, desc in [
        ("meta",          "Metadata: version, fetchDate, prevDate, fetchDateDisplay, prevDateDisplay"),
        ("accounts[]",    "Array of investment accounts. Each contains: id, name, wrapper, provider, holdings[]"),
        ("holdings[]",    "Per-account holdings: id, ticker, name, units, price, prevPrice, priceDate, costBasis, purchaseDate, bucket, fetchStatus, dividendsReceived"),
        ("history[]",     "Daily snapshots: date, totalValue, prices{}, notes. One entry per trading day."),
        ("crypto[]",      "Cryptocurrency positions: id, name, units, price, prevPrice, priceDate"),
        ("cryptoHistory[]", "Daily crypto portfolio total: date, totalValue"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run1 = p.add_run(f"{field}  ")
        run1.font.name = "Courier New"; run1.font.size = Pt(9); run1.font.color.rgb = GOLD
        run2 = p.add_run(desc)
        run2.font.name = "Calibri"; run2.font.size = Pt(9.5); run2.font.color.rgb = LIGHT

    heading2(doc, "Key Data Rules")
    for b in [
        "prevDate / fetchDate: Only rolled forward when a new trading day is first fetched — prevents 'since last update' showing zero on same-day re-runs.",
        "History deduplication: If a history entry already exists for today, it is updated in-place (not appended). No duplicate date entries.",
        "Price currency: Yahoo returns UK securities in pence (GBp/GBX). fetch_prices.py divides by 100 to convert to £.",
        "Crypto prices: Returned directly in GBP — no conversion needed.",
        "Abort on empty fetch: If no prices are retrieved, the script exits with code 1 — the JSON is never overwritten with stale/zero prices.",
    ]:
        bullet(doc, b)

    # ── 6. Security ────────────────────────────────────────────────────────────
    heading1(doc, "6. Security Architecture")

    heading2(doc, "Password Protection")
    for b in [
        "A login screen is displayed before any data or application code is loaded.",
        "Passwords are hashed using SHA-256 (window.crypto.subtle) — the plaintext never leaves the browser.",
        "The stored value is a SHA-256 hash in index.html source — not a plaintext password.",
        "To change the password: echo -n 'newpassword' | shasum -a 256 and replace the hash in index.html.",
        "Limitation: this is client-side only — a user viewing page source could extract the hash and attempt offline cracking. A strong, unique password is recommended.",
    ]:
        bullet(doc, b)

    heading2(doc, "GitHub Secrets")
    for b in [
        "YAHOO_APP_PASSWORD: Yahoo Mail app password for SMTP. Stored as a GitHub repository secret — never in code.",
        "The GitHub Personal Access Token (PAT) used for pushes is embedded in the git remote URL on the local machine only — it is not committed to the repository.",
        "Recommendation: rotate the PAT periodically via GitHub Settings → Developer settings → Personal access tokens.",
    ]:
        bullet(doc, b)

    heading2(doc, "Data Privacy")
    for b in [
        "The repository is public — source code and JSON data files are visible to anyone with the URL.",
        "Portfolio values and holdings data are contained in price_history-2.json which is publicly accessible.",
        "Mitigation: the password gate prevents casual rendering; the public URL is not widely shared.",
        "For higher privacy, upgrading to GitHub Pro ($4/month) allows private repos with Pages support.",
    ]:
        bullet(doc, b)

    # ── 7. Application Tabs ────────────────────────────────────────────────────
    heading1(doc, "7. Application — Tab Reference")

    tabs = [
        ("Dashboard", "Main overview. Shows total portfolio value, since-last-update change, cost basis, total gain, annualised return, asset bucket breakdown, fetch status, and a portfolio value history chart."),
        ("Accounts", "Drill-down into individual accounts. Summary cards per account with value, gain/loss. Click any account to see all holdings with inline price editing capability."),
        ("Allocation", "Asset allocation view. Shows portfolio breakdown by bucket category (e.g. Global Equity, Core Multi-Asset, High Beta/Blockchain) with percentage bars."),
        ("History", "Price history table and chart for all recorded date snapshots. Filter by 24H / 1M / 3M / 6M / 1Y / All."),
        ("Goal", "Retirement goal tracking. Visual projection chart showing Bear / Projected / Bull scenarios toward the £1,400,000 February 2031 goal. Shows estimated dates to reach £1M and full goal. Projection table with value at retirement."),
        ("Analysis", "Displays AI-generated portfolio commentary loaded from analysis.json. Non-blocking — tab functions without the file."),
        ("Crypto", "Cryptocurrency portfolio. Shows BTC, ETH, XRP, ADA positions with current GBP value, 24h change, and crypto portfolio history chart."),
        ("Retirement Planner", "Separate page (retirement.html). Models retirement income scenarios including state pension, drawdown rates, ISA/SIPP balances, and sustainable withdrawal calculations."),
    ]
    for name, desc in tabs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        run1 = p.add_run(f"{name}:  ")
        run1.font.bold = True; run1.font.size = Pt(10.5); run1.font.color.rgb = GOLD; run1.font.name = "Calibri"
        run2 = p.add_run(desc)
        run2.font.size = Pt(10); run2.font.color.rgb = DARKBLUE; run2.font.name = "Calibri"

    # ── 8. Supported Instruments ───────────────────────────────────────────────
    heading1(doc, "8. Supported Instruments & Tickers")
    body(doc, "The following instruments are tracked with automated daily price fetches from Yahoo Finance:")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Ticker", "Name", "Yahoo Symbol", "Type"]):
        set_cell_bg(hdr.cells[i], "0a1420")
        run = hdr.cells[i].paragraphs[0].add_run(h.upper())
        run.font.bold = True; run.font.size = Pt(8); run.font.color.rgb = GOLD; run.font.name = "Calibri"
    instruments = [
        ("BCHS",  "Invesco CoinShares Global Blockchain ETF", "BCHS.L",          "ETF (LSE)"),
        ("SWDA",  "iShares Core MSCI World ETF",               "SWDA.L",          "ETF (LSE)"),
        ("BNY",   "BNY Mellon Multi-Asset Balanced Inst W Acc","0P0000X2GH.L",    "Unit Trust"),
        ("HSBC",  "HSBC Global Strategy Balanced C Acc",       "0P0000WN82.L",    "Unit Trust"),
        ("HISL",  "HSBC Islamic Global Equity Index IC GBP",   "0P0001IVNK.L",    "Unit Trust"),
        ("LION",  "Liontrust Sustainable Future Managed 6 Net","0P0000Y3Y1.L",    "Unit Trust"),
        ("PMIT",  "Premier Miton Diversified Growth D Acc",    "0P0001PKIS.L",    "Unit Trust"),
        ("VLS80", "Vanguard LifeStrategy 80% Equity A Acc",    "0P0000TKZM.L",    "Unit Trust"),
        ("RLSE",  "Royal London Sterling Extra Yield Bond A",  "0P000023IY.L",    "Unit Trust"),
        ("VGAC",  "Vanguard FTSE Global All Cap Index GBP Acc","0P00018XAR.L",    "Unit Trust"),
        ("BTC",   "Bitcoin",                                   "BTC-GBP",         "Cryptocurrency"),
        ("ETH",   "Ethereum",                                  "ETH-GBP",         "Cryptocurrency"),
        ("XRP",   "XRP",                                       "XRP-GBP",         "Cryptocurrency"),
        ("ADA",   "Cardano",                                   "ADA-GBP",         "Cryptocurrency"),
    ]
    for idx, row_data in enumerate(instruments):
        row = tbl.add_row()
        set_row_shading(row, "121e2b" if idx % 2 == 0 else "0f1923")
        for i, val in enumerate(row_data):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = "Courier New" if i in (0, 2) else "Calibri"
            run.font.color.rgb = GOLD if i == 0 else LIGHT

    body(doc, "Note: Unit trusts return prices in pence (GBp) from Yahoo Finance. The fetch script automatically divides by 100 to convert to GBP. ETF prices on .L exchange are also returned in pence.", italic=True, colour=GREY)

    # ── 9. Operations & Maintenance ────────────────────────────────────────────
    heading1(doc, "9. Operations & Maintenance")

    heading2(doc, "Daily Operations (Automated)")
    body(doc, "No manual intervention is required for daily price updates. The GitHub Actions workflow runs automatically at 16:45 UTC on weekdays.")

    heading2(doc, "Manual Tasks")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Task", "How", "Frequency"]):
        set_cell_bg(hdr.cells[i], "0a1420")
        run = hdr.cells[i].paragraphs[0].add_run(h.upper())
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = GOLD; run.font.name = "Calibri"
    tasks = [
        ("Update pension value", "When running fetch manually: python fetch_prices.py 13500 (pass new value as argument)", "Monthly / on statement"),
        ("Sync local file", "Double-click pull_updates.command", "After receiving daily email"),
        ("Change password", "echo -n 'newpw' | shasum -a 256 → update PASSWORD_HASH in index.html → git push", "As required"),
        ("Rotate GitHub PAT", "GitHub Settings → Personal access tokens → regenerate. Update git remote URL locally.", "Annually / on expiry"),
        ("Add new holding", "Edit price_history-2.json — add holding to relevant account, add ticker to YAHOO_SYMBOLS in fetch_prices.py. Git commit + push.", "On portfolio change"),
        ("Update analysis.json", "Replace contents with new AI-generated analysis. Git commit + push.", "Monthly / as needed"),
        ("Monitor Actions", "Check github.com/jwdaniells/portfolio-tracker/actions for failed runs", "If email not received"),
        ("Trigger manual fetch", "GitHub Actions tab → fetch-prices.yml → Run workflow", "Ad hoc / testing"),
    ]
    for idx, row_data in enumerate(tasks):
        row = tbl.add_row()
        set_row_shading(row, "121e2b" if idx % 2 == 0 else "0f1923")
        for i, val in enumerate(row_data):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = LIGHT

    heading2(doc, "Troubleshooting")
    for issue, fix in [
        ("No email received", "Check github.com/.../actions for failed runs. Verify YAHOO_APP_PASSWORD secret is set and Yahoo 'Less secure app access' / app passwords are enabled."),
        ("Prices not updating", "Check Actions log for fetch errors. Yahoo Finance API occasionally returns empty responses — retry is automatic next day."),
        ("Browser shows old data", "Hard refresh (Cmd+Shift+R). Cache-busting is built in via ?v=timestamp on JSON URLs."),
        ("Password forgotten", "Generate a new hash locally and update index.html. Commit and push."),
        ("Git push rejected", "Run git pull --rebase first (GitHub Actions may have committed new prices since your last pull)."),
        ("Fetch script error locally", "Ensure requests is installed: pip install requests. The script self-installs on GitHub Actions."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(4)
        r1 = p.add_run(f"{issue}: ")
        r1.font.bold = True; r1.font.color.rgb = GOLD; r1.font.size = Pt(10); r1.font.name = "Calibri"
        r2 = p.add_run(fix)
        r2.font.color.rgb = DARKBLUE; r2.font.size = Pt(10); r2.font.name = "Calibri"

    # ── 10. Key URLs & References ───────────────────────────────────────────────
    heading1(doc, "10. Key URLs & References")
    refs = [
        ("Live Tracker (GitHub Pages)", "https://jwdaniells.github.io/portfolio-tracker/"),
        ("GitHub Repository",           "https://github.com/jwdaniells/portfolio-tracker"),
        ("GitHub Actions",              "https://github.com/jwdaniells/portfolio-tracker/actions"),
        ("GitHub Secrets",              "https://github.com/jwdaniells/portfolio-tracker/settings/secrets/actions"),
        ("Yahoo Finance (LSE prices)",  "https://query1.finance.yahoo.com/v8/finance/chart/SWDA.L"),
        ("Notification Email",          "jdaniells@yahoo.co.uk"),
    ]
    for label, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        r1 = p.add_run(f"{label}:  ")
        r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = GOLD; r1.font.name = "Calibri"
        r2 = p.add_run(url)
        r2.font.size = Pt(9.5); r2.font.color.rgb = BLUE; r2.font.name = "Courier New"

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 — RETIREMENT PLANNER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PART 2 — RETIREMENT PLANNER")
    run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = GOLD; run.font.name = "Calibri"
    doc.add_paragraph()
    add_horizontal_rule(doc)

    # ── R1. Purpose & Overview ─────────────────────────────────────────────────
    heading1(doc, "1. Purpose & Overview")
    body(doc, (
        "The Retirement Planner is a standalone financial modelling tool that works alongside the "
        "Portfolio Tracker. Its purpose is to project whether the combined investment portfolio will "
        "sustain a target retirement income through to age 95, accounting for UK income tax, state "
        "pension, defined benefit pensions, PCLS lump sums, mortgage obligations, and inflation."
    ))
    heading2(doc, "Key Objectives")
    for b in [
        "Model year-by-year portfolio drawdown across John's and Elaine's accounts (SIPP, ISA) from retirement in 2031.",
        "Optimise the SIPP drawdown split between John and Elaine to minimise combined income tax each year.",
        "Apply UK 2025/26 income tax rules including personal allowance taper above £100k and all three tax bands.",
        "Model three growth scenarios (Bear 3%, Central 5%, Bull 7% real return) simultaneously.",
        "Show portfolio exhaustion age — will the money last to 95?",
        "Calculate PCLS (25% tax-free lump sum) from SIPPs at retirement and model use to clear the mortgage.",
        "Project inheritance tax (IHT) liability including property value, NRB, RNRB, and taper rules.",
        "Allow all assumptions to be edited live in the browser with instant recalculation.",
    ]:
        bullet(doc, b)

    # ── R2. Architecture ───────────────────────────────────────────────────────
    heading1(doc, "2. Architecture & Technology")
    body(doc, (
        "The Retirement Planner is implemented as a React single-page application, identical in "
        "technology to the Portfolio Tracker. It is served from a separate entry point "
        "(retirement.html) which dynamically fetches, Babel-transforms and mounts "
        "retirement_planner.jsx. The planner reads live portfolio values directly from "
        "price_history-2.json so projections always start from current balances."
    ))

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["File", "Role", "Notes"]):
        set_cell_bg(hdr.cells[i], "0a1420")
        run = hdr.cells[i].paragraphs[0].add_run(h.upper())
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = GOLD; run.font.name = "Calibri"
    for idx, (f, r, n) in enumerate([
        ("retirement.html",         "Entry point",    "Fetches + Babel-transforms the JSX; mounts React app; no password gate"),
        ("retirement_planner.jsx",  "Main component", "All logic, UI and modelling in a single JSX file (~700 lines)"),
        ("price_history-2.json",    "Data source",    "Live portfolio values loaded on startup to seed projection balances"),
    ]):
        row = tbl.add_row()
        set_row_shading(row, "121e2b" if idx % 2 == 0 else "0f1923")
        for i, val in enumerate([f, r, n]):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.name = "Courier New" if i == 0 else "Calibri"
            run.font.color.rgb = GOLD if i == 0 else LIGHT

    # ── 3. Projection Model ───────────────────────────────────────────────────
    heading1(doc, "3. Financial Projection Model")

    heading2(doc, "Pre-Retirement Growth Phase (2026–2031)")
    body(doc, (
        "Before the retirement year, the model grows each account forward year-by-year applying "
        "the selected growth rate plus inflation (nominal return). Monthly contributions of £2,011 "
        "are distributed 70% to John's SIPP, 20% to Elaine's SIPP, and 5% each to their ISAs."
    ))

    heading2(doc, "Retirement Drawdown Phase (2031 onwards)")
    for b in [
        "Each year, the model identifies a target net income requirement (base £60,000/year post-mortgage clearing).",
        "Before the mortgage is cleared (age 67), an additional £19,200/year (£1,600/month) is added to the income target.",
        "Stepped income reductions apply in later years: -10% from age 75, -15% from age 80, -20% from age 85.",
        "Guaranteed income (DB pensions, State Pension, LSEG pension) is deducted from the target before any portfolio draw.",
        "An optimisation loop (1% SIPP split increments) finds the exact John/Elaine SIPP draw split that minimises total tax.",
        "ISA withdrawals are used only after SIPP capacity is exhausted — ISA income is tax-free.",
        "Remaining balances grow at the nominal rate each year after drawdown.",
        "The model runs year-by-year until John reaches age 95 or the portfolio is exhausted.",
    ]:
        bullet(doc, b)

    heading2(doc, "PCLS (Pension Commencement Lump Sum)")
    for b in [
        "At retirement (John age 63, 2031), 25% of both SIPPs is taken as tax-free PCLS.",
        "The combined PCLS is used to clear the outstanding mortgage balance (modelled at £250,000 at retirement).",
        "After PCLS is taken, the remaining SIPP balances continue to grow and be drawn for income.",
        "The mortgage cost (£1,600/month) is removed from the income target from the year of clearance onwards.",
    ]:
        bullet(doc, b)

    heading2(doc, "Income Sources Modelled")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Source", "Owner", "Annual Amount / Notes"]):
        set_cell_bg(hdr.cells[i], "0a1420")
        run = hdr.cells[i].paragraphs[0].add_run(h.upper())
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = GOLD; run.font.name = "Calibri"
    for idx, (src, owner, notes) in enumerate([
        ("State Pension",        "John",   "£11,502/year from age 67 (2025/26 full new state pension)"),
        ("State Pension",        "Elaine", "£11,502/year from age 67"),
        ("Atkins DB Pension",    "John",   "£5,000/year from age 65 — not inflation-indexed"),
        ("Pfizer DB Pension",    "John",   "£2,500/year from age 65 — not inflation-indexed"),
        ("Disney DB Pension",    "Elaine", "£1,237/year from age 65 — not inflation-indexed"),
        ("Nippon Life Pension",  "Elaine", "£135/year from age 65 — not inflation-indexed"),
        ("LSEG DC Pension",      "John",   "Modelled as 5% annuity equivalent from retirement value"),
        ("John SIPP drawdown",   "John",   "Tax-assessed at marginal rate; optimised split with Elaine"),
        ("Elaine SIPP drawdown", "Elaine", "Tax-assessed at marginal rate; receives portion of combined draw"),
        ("ISA withdrawals",      "Both",   "Tax-free; drawn after SIPP capacity used; split proportionally"),
    ]):
        row = tbl.add_row()
        set_row_shading(row, "121e2b" if idx % 2 == 0 else "0f1923")
        for i, val in enumerate([src, owner, notes]):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = LIGHT

    # ── R4. Tax Engine ─────────────────────────────────────────────────────────
    heading1(doc, "4. UK Tax Engine")
    body(doc, (
        "The planner contains a built-in UK income tax engine using 2025/26 rates with "
        "frozen thresholds. It is applied independently to John and Elaine to calculate "
        "their net income after tax each year."
    ))
    for b in [
        "Personal Allowance: £12,570 (tapered by £1 for every £2 of income above £100,000)",
        "Basic rate: 20% on income between Personal Allowance and £50,270.",
        "Higher rate: 40% on income between £50,271 and £125,140.",
        "Additional rate: 45% on income above £125,140.",
        "ISA withdrawals are excluded from the tax calculation entirely.",
        "The optimiser minimises combined tax by adjusting the relative SIPP draw between John and Elaine.",
        "A binary-search algorithm iterates to find the exact gross draw that meets the target net income.",
    ]:
        bullet(doc, b)
    body(doc, (
        "Example: if John has £30,000 of guaranteed income, his SIPP drawdown is limited to "
        "£20,270 before hitting the higher-rate band. Excess is then drawn from Elaine's SIPP "
        "or the ISA to minimise total tax."
    ), italic=True, colour=GREY)

    # ── R5. IHT Projection ─────────────────────────────────────────────────────
    heading1(doc, "5. Inheritance Tax Projection")
    body(doc, (
        "A separate IHT tab projects the potential inheritance tax liability on the combined "
        "estate (investment portfolio + property) across the three growth scenarios, year by year "
        "until John's age 95."
    ))
    for b in [
        "Property value projected at 3% nominal growth per year from current £1.2M.",
        "Remaining mortgage (currently ~£350,000) reduces estate value until cleared.",
        "Nil-Rate Band (NRB): £325,000 per person — £650,000 combined.",
        "Residence Nil-Rate Band (RNRB): £175,000 per person — £350,000 combined — for direct descendant inheritance.",
        "RNRB tapers by £1 per £2 of estate value above £2,000,000.",
        "IHT is charged at 40% on estate value above the combined threshold.",
        "SIPPs are excluded from the taxable estate before 6 April 2027 (current rules). From 2027, defined contribution pension pots enter the estate (Budget 2024 change) — the model reflects this automatically.",
        "ISAs are included in the taxable estate.",
    ]:
        bullet(doc, b)

    # ── R6. Sections / UI ──────────────────────────────────────────────────────
    heading1(doc, "6. Application Sections")
    body(doc, "The planner is divided into five sections accessible via a tab bar:")
    sections = [
        ("Overview",
         "Summary dashboard. Shows total portfolio at retirement (Bear/Central/Bull), estimated exhaustion age "
         "across scenarios, portfolio value at key ages (80, 90), and the year the portfolio is projected to "
         "run out (if applicable). A colour-coded chart shows portfolio trajectories across all three scenarios."),
        ("Year by Year",
         "Detailed annual projection table. Columns: Year, Ages, Income Target, Guaranteed Income, Portfolio Draw, "
         "Tax Paid, Net Income, Total Portfolio, per-account balances. Scrollable; highlights exhaustion year in red."),
        ("Income Analysis",
         "Breakdown of income sources for each year: DB pensions by name, State Pension, SIPP drawdown (John/Elaine), "
         "ISA draw, and total tax. Useful for understanding tax efficiency over time."),
        ("IHT Planner",
         "Inheritance tax projection showing estate value (portfolio + property − mortgage), combined NRB/RNRB, "
         "taxable amount, and estimated IHT bill across all three scenarios."),
        ("Assumptions",
         "Editable form for all model inputs. Changes apply immediately — no save/submit required. "
         "Includes: portfolio values (auto-loaded from price_history-2.json), contribution amounts and split, "
         "growth rate scenarios, target income, DB pension details, state pension ages, income step-down ages, "
         "mortgage details, PCLS options, and IHT property assumptions."),
    ]
    for name, desc in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        run1 = p.add_run(f"{name}:  ")
        run1.font.bold = True; run1.font.size = Pt(10.5); run1.font.color.rgb = GOLD; run1.font.name = "Calibri"
        run2 = p.add_run(desc)
        run2.font.size = Pt(10); run2.font.color.rgb = DARKBLUE; run2.font.name = "Calibri"

    # ── R7. Key Assumptions & Defaults ────────────────────────────────────────
    heading1(doc, "7. Default Assumptions")
    body(doc, "All values are editable in the Assumptions tab. The initial defaults are:")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Assumption", "Default Value", "Notes"]):
        set_cell_bg(hdr.cells[i], "0a1420")
        run = hdr.cells[i].paragraphs[0].add_run(h.upper())
        run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = GOLD; run.font.name = "Calibri"
    for idx, (a, v, n) in enumerate([
        ("Retirement year",          "2031",           "March 2031, John age 63"),
        ("Target net income",        "£60,000/year",   "Post-mortgage, post-tax, in today's money"),
        ("Monthly contributions",    "£2,011",         "Combined; 70% John SIPP / 20% Elaine SIPP / 10% ISAs"),
        ("Growth — Bear",            "3% real",        "After inflation"),
        ("Growth — Central",         "5% real",        "After inflation (default view)"),
        ("Growth — Bull",            "7% real",        "After inflation"),
        ("Inflation rate",           "2.5%",           "Used to convert real to nominal returns"),
        ("Longevity",                "Age 95",         "John's plan-to age"),
        ("State pension",            "£11,502/year",   "Full new state pension (2025/26); from age 67"),
        ("Mortgage balance",         "£250,000",       "At retirement; cleared via PCLS"),
        ("Mortgage cost",            "£1,600/month",   "Added to income target until cleared"),
        ("House value (IHT)",        "£1,200,000",     "Current market value; grown at 3%/year"),
        ("PCLS use",                 "Clear mortgage", "25% tax-free lump from both SIPPs applied to mortgage"),
        ("Income step-down age 75",  "-10%",           "Reflects lower activity and travel spend"),
        ("Income step-down age 80",  "-15%",           "More home-based lifestyle"),
        ("Income step-down age 85",  "-20%",           "Lower discretionary spend"),
    ]):
        row = tbl.add_row()
        set_row_shading(row, "121e2b" if idx % 2 == 0 else "0f1923")
        for i, val in enumerate([a, v, n]):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = LIGHT

    # ── 8. Limitations & Caveats ─────────────────────────────────────────────
    heading1(doc, "8. Limitations & Caveats")
    for b in [
        "Tax rules are based on 2025/26 UK rates with frozen thresholds — future governments may change these significantly.",
        "SIPP legislation change (Budget 2024): From April 2027, SIPPs enter the taxable estate. The model already reflects this — SIPPs are excluded from the estate before 2027 and included from 2027 onwards. A dashed vertical marker on the IHT charts shows the transition point.",
        "Growth rates are assumed constant per year — actual returns will be volatile.",
        "Inflation is modelled as a flat 2.5% — actual inflation will vary.",
        "DB pension amounts are held constant (not inflation-linked) — some schemes may provide inflation uplift.",
        "The model does not currently account for annuity purchase, care home costs, or inter-vivos gifting strategies.",
        "Property value projections are illustrative — regional and market variations apply.",
        "This tool is for personal planning purposes only and does not constitute regulated financial advice.",
    ]:
        bullet(doc, b)

    # ── Footer note ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Daniells Portfolio & Retirement Tracker — System Documentation v1.20 — {TODAY} — Personal & Confidential")
    run.font.size = Pt(8); run.font.color.rgb = GREY; run.font.name = "Calibri"

    # ── Save ───────────────────────────────────────────────────────────────────
    out = DIR / "Portfolio_Tracker_System_Documentation.docx"
    doc.save(out)
    print(f"✓ Saved: {out}")
    return out


if __name__ == "__main__":
    build_doc()
